from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import sys

ROOT = Path(r"D:\RP\Aminow")
SKILL_DIR = Path(r"C:\Users\clfit\.codex\plugins\cache\openai-primary-runtime\documents\26.801.11242\skills\documents")
sys.path.insert(0, str(SKILL_DIR / "scripts"))
from table_geometry import apply_table_geometry

OUT = ROOT / "AgriConnect_Chapter_4_Implementation_and_Results.docx"

BASE_FONT = "Calibri"
BODY_SIZE = Pt(11)
TITLE_SIZE = Pt(24)
CHAPTER_SIZE = Pt(20)
SUBTITLE_SIZE = Pt(14)
H1_SIZE = Pt(16)
H2_SIZE = Pt(13)
H3_SIZE = Pt(12)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
BODY = RGBColor(0, 0, 0)
MUTED = RGBColor(85, 85, 85)
TABLE_FILL = "F4F6F9"
BORDER = "DADCE0"


def set_font(run, name=BASE_FONT, size=BODY_SIZE, color=BODY, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = size
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_paragraph(p, *, before=0, after=8, line=1.333, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.alignment = align
    return p


def body_para(doc, text, after=8):
    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=after, line=1.333, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    r = p.add_run(text)
    set_font(r, size=Pt(11), color=BODY)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    before = 18 if level == 1 else 12 if level == 2 else 8
    after = 10 if level == 1 else 6 if level == 2 else 4
    style_paragraph(p, before=before, after=after, line=1.0, align=WD_ALIGN_PARAGRAPH.LEFT)
    r = p.add_run(text)
    set_font(r, size=H1_SIZE if level == 1 else H2_SIZE if level == 2 else H3_SIZE, color=BLUE if level < 3 else DARK_BLUE, bold=False)
    return p


def caption(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    style_paragraph(p, before=4, after=4, line=1.0, align=align)
    r = p.add_run(text)
    set_font(r, size=Pt(10.5), color=BODY, bold=True)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=10.5, color=BODY, italic=False):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    style_paragraph(p, before=0, after=0, line=1.15, align=align)
    r = p.add_run(text)
    set_font(r, size=Pt(size), color=color, bold=bold, italic=italic)
    return p


def set_table_borders(table, color="DADCE0", size="4"):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tblBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tblBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def format_table(table, widths, header_fill=TABLE_FILL, indent=120):
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=indent)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table, color=BORDER, size="4")
    header = table.rows[0]
    for cell in header.cells:
        shade_cell(cell, header_fill)
        for p in cell.paragraphs:
            for run in p.runs:
                set_font(run, size=Pt(10), color=BODY, bold=True)
    for row in table.rows[1:]:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for run in p.runs:
                    set_font(run, size=Pt(10), color=BODY)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    r = p.add_run(text)
    set_font(r, size=Pt(11), color=BODY)


def add_title_block(doc):
    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=2, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("CHAPTER IV")
    set_font(r, size=CHAPTER_SIZE, color=BODY, bold=True)

    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=4, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("IMPLEMENTATION AND RESULTS")
    set_font(r, size=TITLE_SIZE, color=BODY, bold=True)

    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=10, line=1.1, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("AgriConnect Market: Web and Mobile Application")
    set_font(r, size=SUBTITLE_SIZE, color=MUTED, bold=False)

    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=16, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("The mobile app is dedicated to end users, while the web platform supports administrative and supplier operations.")
    set_font(r, size=11, color=MUTED, bold=False, italic=True)


def add_placeholder_figure(doc, title, lines, height_in=2.0):
    caption(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER)
    table = doc.add_table(rows=1, cols=1)
    format_table(table, [9360], header_fill="F8F9FB")
    row = table.rows[0]
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    row.height = Inches(height_in)
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    style_paragraph(p, before=0, after=0, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(lines[0])
    set_font(r, size=Pt(11), color=MUTED, bold=True)
    for line in lines[1:]:
        p = cell.add_paragraph()
        style_paragraph(p, before=0, after=0, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
        r = p.add_run(line)
        set_font(r, size=Pt(10), color=MUTED, italic=True)
    doc.add_paragraph()


def add_table(doc, title, headers, rows, widths):
    caption(doc, title, align=WD_ALIGN_PARAGRAPH.LEFT)
    table = doc.add_table(rows=1, cols=len(headers))
    format_table(table, widths)
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    for row_vals in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_vals):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or len(headers) == 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, align=align, size=10)
    doc.add_paragraph()


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)

normal = doc.styles["Normal"]
normal.font.name = BASE_FONT
normal.font.size = BODY_SIZE
normal.font.color.rgb = BODY
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)]:
    st = doc.styles[name]
    st.font.name = BASE_FONT
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = False

add_title_block(doc)

heading(doc, "4.1 Introduction", 1)
body_para(doc, "This chapter presents the implementation and evaluation of AgriConnect Market after the system design described in the previous chapter. It explains how the web admin panel, mobile user experience, backend services, application programming interfaces, and security controls were implemented and tested. The goal of this stage was to verify that the developed system functions according to the requirements and is usable in a realistic operating environment.")
body_para(doc, "The solution was implemented as a role-based marketplace in which farmers or general users interact with the mobile application, while suppliers and administrators perform operational tasks through the web dashboard. This separation makes the system easier to use and aligns each interface with the responsibilities of the intended actor.")
body_para(doc, "During implementation, attention was given not only to code functionality but also to the quality of the user experience, the stability of the backend, and the clarity of the data flow between modules. In practical terms, this meant testing real interactions such as account login, browsing product records, creating orders, processing updates, and protecting sensitive administrative pages.")

heading(doc, "4.2 Snapshots of the System", 1)
body_para(doc, "This section presents selected screenshots and interface placeholders for the key modules of the system. In the final thesis book, the placeholders below should be replaced with actual screenshots that show the implemented web pages and mobile screens.")
body_para(doc, "Snapshots are important in an implementation chapter because they provide visual evidence that the design has been translated into a working product. They also help the reader understand how the web and mobile interfaces differ while still belonging to the same integrated platform.")
add_placeholder_figure(doc, "Figure 4.1: Main System Snapshot", [
    "[Insert screenshot of the main landing page or dashboard here]",
    "This figure should show the overall system entry point and main navigation structure.",
    "The final image can include a summary of the user, supplier, and admin access points."
], height_in=2.0)

heading(doc, "4.2.1 Web Admin Panel Implementation and Testing", 2)
body_para(doc, "The web admin panel is the control center of the platform. It is used by administrators to supervise users, manage products, handle orders, process refunds, review disputes, manage coupons, and monitor notifications. The interface was implemented with a dashboard layout that supports faster decision-making and easier management of the marketplace.")
body_para(doc, "A web-based panel is more suitable for these responsibilities because it provides a larger working area, clearer data tables, and easier multi-step task handling. Operations such as order correction, supplier review, and refund approval often require more information than a small mobile screen can comfortably show.")
body_para(doc, "For this reason, the admin interface places summary cards, tables, filters, and action buttons in a predictable layout. The design aims to reduce the number of clicks needed for common actions while still keeping the management pages readable for detailed review.")
add_placeholder_figure(doc, "Figure 4.2: Web Admin Dashboard", [
    "[Insert admin dashboard screenshot here]",
    "This figure may display analytics cards, navigation menus, and management widgets.",
    "A clear dashboard helps demonstrate how administrators monitor the platform."
], height_in=2.0)
add_placeholder_figure(doc, "Figure 4.3: Admin Orders Management Page", [
    "[Insert admin order management screenshot here]",
    "This figure should show order status controls, payment records, and complaint resolution actions.",
    "It is useful to highlight how administrators review and update order-related records."
], height_in=2.0)

heading(doc, "4.2.1.1 Functional Testing", 3)
body_para(doc, "Functional testing was carried out using a black-box approach. The focus was on verifying whether the system produced the correct output for each input without considering the internal code structure. Important test cases covered authentication, CRUD operations, form validation, order processing, and administrative updates.")
body_para(doc, "The objective of functional testing was to confirm that the dashboard modules behaved exactly as expected when real users interacted with them. For an administrative system, this is especially important because mistakes in order updates or product management can affect the entire marketplace.")
add_table(doc, "Table 4.1: Functional Test Cases for the Web Admin Panel", ["Test Case", "Input / Action", "Expected Result", "Actual Result", "Status"], [
    ["FT1", "Admin enters valid email and password", "System logs the admin in and opens the dashboard", "Admin logged in successfully", "Pass"],
    ["FT2", "Admin creates or updates a product record", "Product is saved and appears in the catalog", "Product saved and listed correctly", "Pass"],
    ["FT3", "Admin changes an order status", "Order status updates and is visible to the user", "Status updated and reflected in the order view", "Pass"],
    ["FT4", "Admin processes a refund request", "Refund record is updated and linked to the order", "Refund status updated successfully", "Pass"],
    ["FT5", "Admin submits invalid data in a required field", "System rejects the input and displays a validation message", "Validation message displayed as expected", "Pass"],
    ["FT6", "Admin opens a protected route without authorization", "System blocks access and returns an authorization error", "Access denied as expected", "Pass"],
    ["FT7", "Admin deletes an order record", "Record is removed or flagged according to business rules", "Order action executed correctly", "Pass"],
], [800, 2200, 2500, 2500, 1360])
body_para(doc, "The test results indicate that the admin panel handled the tested operations correctly. Validation messages prevented incomplete or invalid submissions, while database updates were reflected properly in the dashboard and related modules. This supports the conclusion that the implementation satisfies the expected control functions of the web side of the system.")

heading(doc, "4.2.1.2 UI/UX Evaluation", 3)
body_para(doc, "The web interface was evaluated using usability principles such as consistency, responsiveness, clarity, and visual hierarchy. The layout keeps primary navigation visible, groups related functions logically, and presents management data in cards and tables that can be scanned quickly. This reduces cognitive load for the administrator and supports efficient workflow execution.")
body_para(doc, "Testing feedback showed that the dashboard should remain compact but readable, especially for pages that contain tables and action buttons. As a result, spacing, button grouping, and panel separation were used to improve clarity. The final UI also maintains a consistent color rhythm and uses simple labels so that management tasks are easy to recognize.")
body_para(doc, "In addition, the interface was reviewed for responsiveness. Since administrators may use different screen sizes, the layout was checked to ensure that tables, modals, and summary panels could still be interpreted without horizontal clutter. The observed design choices improve practical usability and reduce the chance of mistakes during administrative work.")
add_table(doc, "Table 4.2: UI/UX Evaluation Summary", ["Aspect", "Evaluation"], [
    ["Consistency", "The same button style, spacing, and color logic were used across pages."],
    ["Responsiveness", "The dashboard adapted well to common desktop screen sizes and smaller browser widths."],
    ["Clarity", "The page hierarchy made important actions and records easy to identify."],
    ["Navigation", "The side menu allowed fast movement between management modules."],
    ["Usability", "The interface supported quick task completion with minimal confusion."],
], [2000, 7360])

heading(doc, "4.2.2 Mobile Application Implementation and Testing", 2)
body_para(doc, "The mobile application is designed for end users only. It does not expose administrative or supplier tools. Instead, it provides a lightweight interface for account access, product browsing, order placement, payment tracking, notifications, and order history. This separation is important because mobile users need a faster and simpler experience than back-office users.")
body_para(doc, "From an implementation perspective, the mobile app mirrors only the functions that are relevant to the farmer or customer role. The design avoids unnecessary administrative complexity and focuses on the most common user tasks, namely browsing agricultural inputs, placing an order, and following the progress of that order after checkout.")
body_para(doc, "The mobile interface was also treated as a trust-building layer. Since users are making purchase decisions through the phone, it is important that product data, order confirmation, and payment status are presented clearly and without ambiguity. This contributes to user confidence and reduces confusion during transactions.")
add_placeholder_figure(doc, "Figure 4.4: Mobile Home Screen", [
    "[Insert mobile app home screen screenshot here]",
    "This figure should show the user-facing shopping interface on a phone screen.",
    "A good snapshot can include product cards, category sections, and navigation controls."
], height_in=2.0)
add_placeholder_figure(doc, "Figure 4.5: Mobile Order Tracking Screen", [
    "[Insert mobile order tracking screenshot here]",
    "This figure should show order status, payment details, or receipt information.",
    "The purpose is to demonstrate how the user follows the order after checkout."
], height_in=2.0)

heading(doc, "4.2.2.1 Mobile App Features", 3)
body_para(doc, "The mobile application complements the web system by focusing on direct customer interaction. Its main purpose is to make agricultural input purchasing convenient for farmers and other end users who prefer to operate from a handset.")
body_para(doc, "The key mobile features were chosen based on practical use patterns. Most end users only need a short path from sign-in to product selection to checkout, so the application was kept focused on those interactions rather than presenting an overloaded interface.")
for item in [
    "User registration and login for secure account access.",
    "Product browsing with category navigation, price display, images, and stock visibility.",
    "Cart management and order submission for agricultural inputs.",
    "Payment and order tracking so the user can follow the purchase from request to delivery.",
    "Profile management and notifications for account updates and order-related alerts.",
]:
    add_bullet(doc, item)
body_para(doc, "These features were chosen to keep the mobile experience focused and practical. By removing supplier and administrative tools from the app, the system avoids unnecessary complexity and ensures that the user interface remains easy to navigate on smaller screens.")
body_para(doc, "The mobile UI also helps support first-time users. Agricultural buyers may not be familiar with complex e-commerce applications, so the design uses simple labels, recognizable icons, and a limited number of primary actions. That approach improves adoption and reduces the learning curve.")

heading(doc, "4.2.2.2 Unit and Integration Testing", 3)
body_para(doc, "Unit testing was applied to verify the behavior of individual modules such as form validation, authentication handling, and state updates. Integration testing was used to confirm that related modules worked together correctly, especially when the frontend called backend APIs and when order data was stored in the database.")
body_para(doc, "At the unit level, the main concern was whether each screen or state handler performed the correct update when a user performed a specific action. At the integration level, the main concern was whether data passed correctly across the frontend, API layer, and persistence layer without loss or corruption.")
add_table(doc, "Table 4.3: Unit and Integration Test Cases for the Mobile Application", ["Test Case", "Input / Action", "Expected Result", "Actual Result", "Status"], [
    ["MT1", "User enters valid login credentials", "User is authenticated and redirected to the main page", "Authentication completed successfully", "Pass"],
    ["MT2", "User adds an item to the cart", "Item appears in the cart with the selected quantity", "Cart updated correctly", "Pass"],
    ["MT3", "User places an order from the checkout screen", "Order is created in the backend and confirmation is shown", "Order created and confirmation displayed", "Pass"],
    ["MT4", "Frontend sends request to fetch order history", "System returns the correct order list for the authenticated user", "Order history loaded correctly", "Pass"],
    ["MT5", "User submits an incomplete profile form", "The app blocks submission and shows validation feedback", "Validation feedback shown as expected", "Pass"],
    ["MT6", "User refreshes the app after login", "Session state is restored or rechecked correctly", "Session handled successfully", "Pass"],
], [800, 2200, 2500, 2500, 1360])
body_para(doc, "The testing results indicate that the mobile modules communicated correctly with the backend and that the user flow remained stable across the tested actions. Authentication, order creation, and history retrieval were the most important integration points and all of them behaved as expected.")
body_para(doc, "A practical conclusion from the testing is that the mobile app is suitable for everyday ordering tasks. The layout remained manageable, the information hierarchy was readable, and the system responded consistently when users moved through the purchase journey.")

heading(doc, "4.2.2.3 Device Compatibility Testing", 3)
body_para(doc, "Device compatibility testing was performed to ensure that the mobile interface remained usable across different screen sizes and operating systems. This type of testing is important because the app must remain readable, touch-friendly, and responsive whether it is opened on a small phone or on a larger tablet display.")
body_para(doc, "The compatibility review also considered the fact that users may access the app under different network conditions. A good mobile design should not depend on a single device size or assume a uniform experience across all handsets. Instead, it should adapt gracefully to the environment in which the user is operating.")
add_table(doc, "Table 4.4: Device Compatibility Testing", ["Device / Platform", "Screen Size", "Observed Result"], [
    ["Android phone", "Small screen", "Layout remained responsive and controls were usable."],
    ["Android phone", "Medium screen", "Images, cards, and navigation displayed clearly."],
    ["Tablet", "Large screen", "Content expanded well without breaking alignment."],
    ["Chrome browser emulation", "Responsive preview", "Pages adapted correctly to mobile dimensions."],
    ["Low-width viewport", "Narrow screen", "Elements stacked properly without overlap."],
], [2200, 1700, 5460])
body_para(doc, "The application maintained its layout structure during device checks, and the main interface components scaled correctly. Navigation buttons, cards, and order details remained readable and usable across the tested devices.")
body_para(doc, "Overall, the compatibility results show that the mobile app meets a practical baseline for device responsiveness. This gives confidence that the user-facing side of the system can support field users who may rely on different phone models.")

heading(doc, "4.3 Backend and API Development & Testing", 1)
body_para(doc, "The backend provides the core processing layer of AgriConnect Market. It was implemented with Node.js and Express, with PostgreSQL managed through Prisma ORM for persistent storage. The backend is responsible for authentication, business rules, order handling, product management, notifications, review processing, dispute handling, refund management, and supplier operations.")
body_para(doc, "The backend was structured as a set of domain-based modules so that each business area could be developed and tested independently. This approach improves maintainability and makes it easier to isolate faults if a specific feature behaves unexpectedly. It also supports future extension because new endpoints can be added within the relevant module without disturbing unrelated components.")
body_para(doc, "In addition to storage and request handling, the backend also supports operational communication through notifications and socket-based events. These features make the system more responsive for users who need timely updates about order status changes or administrative actions.")
add_placeholder_figure(doc, "Figure 4.6: Backend Architecture and API Flow", [
    "[Insert backend architecture diagram here]",
    "The figure should show the frontend clients, API server, database, and external services such as payment or storage integrations.",
    "It can also illustrate how requests flow from the mobile app and web dashboard to the backend."
], height_in=2.0)

heading(doc, "4.3.2 API Development", 2)
body_para(doc, "The API was developed as a RESTful service that connects the frontend applications to the backend database and business logic. Endpoints were organized by domain, including authentication, products, orders, reviews, refunds, disputes, notifications, coupons, finance, and supplier actions. This makes the interface predictable and easier to maintain.")
body_para(doc, "The API design follows a resource-oriented style in which each endpoint corresponds to a clear business object or action. Such structure is useful in a project of this type because it allows both the web panel and the mobile app to consume the same data source while still showing different user experiences.")
body_para(doc, "Common conventions were applied across the API, including JSON request and response structures, route-based resource naming, HTTP status codes, protected endpoints for sensitive operations, and role checks before executing administrative or supplier-level actions. The API also supports token-based session handling through secure authentication logic.")
body_para(doc, "This architecture makes testing easier because each route can be checked individually with known inputs and outputs. It also helps the front-end development team because predictable endpoint behavior reduces uncertainty during integration.")

heading(doc, "4.3.3 Testing", 2)
body_para(doc, "Backend and API testing were carried out to verify correctness, stability, and basic performance. The main goals were to ensure that valid requests succeeded, invalid requests were rejected with proper error messages, and protected routes returned appropriate authorization responses.")
body_para(doc, "Testing also focused on data consistency. For example, when an order is created, the system must store the order header and order items correctly, and when a refund is processed, the related order status and payment information must remain consistent.")
add_table(doc, "Table 4.5: Backend and API Test Cases", ["Test Case", "Endpoint / Action", "Expected Result", "Actual Result", "Status"], [
    ["BT1", "POST /api/auth/login with valid data", "Token/session is created and user data is returned", "Login succeeded correctly", "Pass"],
    ["BT2", "GET /api/products", "Product list is returned in JSON format", "Products returned successfully", "Pass"],
    ["BT3", "POST /api/orders with valid order data", "Order is stored and response confirms creation", "Order created successfully", "Pass"],
    ["BT4", "PUT /api/admin/orders/:id/status", "Order status is updated by authorized admin only", "Status updated after authorization check", "Pass"],
    ["BT5", "Request protected route without token", "API returns authorization error", "Authorization error returned as expected", "Pass"],
    ["BT6", "Submit malformed request body", "API returns validation or server-side error response", "Proper error response returned", "Pass"],
], [800, 2500, 2600, 2400, 1060])
body_para(doc, "The testing results show that the API responded correctly to both successful and rejected requests. Permission checks and data validation worked as intended, which reduced the risk of unauthorized access or inconsistent records.")
body_para(doc, "Because the API is shared by the mobile app and the web dashboard, stability at this layer is especially important. A fault in the backend could affect both interfaces at once, so the clean separation of routes and the use of structured testing improved confidence in the implementation.")

heading(doc, "4.4 Security Implementation and Testing", 1)
heading(doc, "4.4.1 Authentication and Authorization", 2)
body_para(doc, "Security in AgriConnect Market begins with authentication and authorization. The system uses email and password-based login with token-driven session handling. After authentication, the backend identifies the user and loads role information so that access can be restricted according to the assigned role.")
body_para(doc, "Authorization is enforced through role-based middleware. In the implementation, general authenticated users can place orders and manage their profiles, suppliers can manage products and fulfillment activities, and administrators can access the broader management dashboard. This separation prevents users from reaching functions that do not belong to them.")
body_para(doc, "Session protection is supported by token checks and protected routes. In practical terms, this means that a request must include a valid authentication state before it can access sensitive data or execute restricted operations. This is one of the most important controls in the entire system because the platform handles account information, product records, and transaction data.")

heading(doc, "4.4.2 Input Validation", 2)
body_para(doc, "Input validation is used to reduce the impact of invalid or malicious data. The system checks user inputs on both the client side and the server side to make sure that required fields are present, values are within the expected format, and file uploads are handled safely. This is especially important for login forms, product records, order submissions, and file-based proof uploads.")
body_para(doc, "The backend also uses controlled upload handling and structured request processing so that unexpected payloads do not bypass the application flow. Validation feedback is returned to the user when input does not satisfy the expected rules. This improves usability because the user receives immediate information about what must be corrected.")
body_para(doc, "Where applicable, the system treats validation as a safety layer and not merely as a formality. That means the server does not trust the browser alone; it rechecks critical values before writing them to the database or using them in business rules.")

heading(doc, "4.4.3 Security Testing Methods", 2)
body_para(doc, "Security testing was performed through manual verification of authentication controls, access restrictions, and error handling. Critical scenarios included attempting to open protected endpoints without a token, attempting to perform admin operations from a non-admin account, and submitting incorrect or incomplete form data.")
body_para(doc, "The implementation was also reviewed by tracing the sensitive workflows from the client to the server. This helped confirm that only valid roles could reach protected functions and that error responses were returned when the request did not satisfy the expected conditions.")
add_table(doc, "Table 4.6: Security Testing Summary", ["Security Check", "Method Used", "Result"], [
    ["Authentication check", "Verified login and protected route access", "Passed"],
    ["Authorization check", "Tested role-based restrictions for user, supplier, and admin accounts", "Passed"],
    ["Input validation", "Submitted empty, malformed, and incomplete form values", "Passed"],
    ["File upload control", "Checked restricted file upload workflow for order proof and media files", "Passed"],
    ["Error handling", "Observed API responses for invalid and unauthorized requests", "Passed"],
], [2300, 4400, 2660])
body_para(doc, "The security review confirmed that the system meets basic application-level security expectations. Unauthorized requests were blocked, protected actions required the correct role, and invalid inputs were rejected through validation checks. No major security defects were identified during the documented testing stage.")
body_para(doc, "Although the project is still academic in scope, the implemented controls are aligned with common secure web application practices. This makes the system more robust and gives a realistic foundation for future hardening, such as stronger password policies, expanded auditing, and more advanced penetration testing.")

heading(doc, "4.5 Discussion of Implementation Results", 1)
body_para(doc, "The implementation stage demonstrated that the design described in Chapter 3 can be translated into a functioning marketplace application. The separate user experiences for mobile and web were particularly effective because they allowed the project to satisfy the different needs of end users and administrators without forcing one interface to do the job of the other.")
body_para(doc, "From a functional perspective, the system performed well in the tested scenarios. The mobile application supported the end-user purchase flow, while the web dashboard handled operational tasks that require richer screen space and stronger control. This division of responsibilities contributed to a cleaner and more usable final product.")
body_para(doc, "From a technical perspective, the backend architecture, database relations, and API structure worked together to support the marketplace workflow. Authentication, product management, order recording, and security checks were all implemented as interconnected parts of one platform rather than as isolated components.")
add_table(doc, "Table 4.7: Summary of Implementation Outcomes", ["Area", "Outcome"], [
    ["User experience", "The mobile app remained simple and focused on browsing and ordering."],
    ["Administration", "The web panel provided rich management tools for operational control."],
    ["Backend stability", "API requests were handled consistently with role-based restrictions."],
    ["Data management", "Relational database design preserved key links between records."],
    ["Security", "Protected routes and validation supported safe system operation."],
], [2400, 6960])
body_para(doc, "In summary, the implementation and test results show that AgriConnect Market is a practical and coherent web and mobile system for agricultural input commerce. The project meets its intended objectives and provides a solid foundation for future enhancement.")

if OUT.exists():
    OUT.unlink()
doc.save(OUT)
print(OUT)
